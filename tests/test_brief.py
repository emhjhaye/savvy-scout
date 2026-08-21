from pypdf import PdfReader
from docx import Document

from savvy_scout.escalation.brief import (
    build_brief,
    build_capture_brief,
    build_original_notice_pdf,
    record_brief,
)
from savvy_scout.models.notice import Notice
from savvy_scout.sources.ocds_parser import ParsedNotice
from savvy_scout.sweep.dedupe import upsert_notice
from savvy_scout.triage.gates import triage_notice
from savvy_scout.triage.scope_read import save_scope_read
from savvy_scout.escalation.word_documents import (
    build_capture_brief_docx,
    build_internal_addendum_docx,
)

SECTION_TITLES = [
    "TRIAGE SUMMARY",
    "CAPABILITY MAPPING",
    "BLOCKERS & RISKS",
    "DIRECT ASKS",
    "RECOMMENDATION",
]

CAPTURE_SECTION_TITLES = [
    "1. OPPORTUNITY SUMMARY",
    "2. BUYER",
    "3. VALUE",
    "4. ROUTE TO MARKET",
    "5. GATE OUTCOMES",
    "6. PROVISIONAL RATINGS WITH REASONING",
    "7. COMPETITOR PICTURE",
    "8. RISKS",
    "9. OPEN QUESTIONS",
    "10. RECOMMENDED NEXT ACTION",
]


def _pdf_text(path: str) -> str:
    reader = PdfReader(path)
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def _docx_text(path: str) -> str:
    document = Document(path)
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    for section in document.sections:
        parts.extend(paragraph.text for paragraph in section.footer.paragraphs)
    return "\n".join(parts)


def _make_notice(conn):
    # Matches both Fintech ("bank" + "payments platform" coupling) and Energy
    # ("energy" + "smart grid" coupling) -> contested -> Gate 1 FLAG.
    notice = Notice(
        ref="REF-BRIEF-1",
        title="Ambiguous Sector Notice For Brief Test",
        buyer="Some Bank",
        source="Find a Tender",
        notice_type="UK3",
        uk_stage="UK3",
        raw_json="{}",
        notice_url="https://www.find-tender.service.gov.uk/Notice/REF-BRIEF-1",
    )
    parsed = ParsedNotice(
        notice=notice,
        text_blob="real-time payments platform integration with smart grid billing systems "
        "for our energy trading desk",
        tender_status="active",
    )
    notice_id = upsert_notice(conn, parsed)
    triage_notice(conn, notice_id)  # gate1 contested -> FLAG
    return notice_id


def test_build_brief_has_all_sections_and_warning(conn, tmp_path):
    notice_id = _make_notice(conn)
    output_dir = tmp_path / "briefs"
    path = build_brief(conn, notice_id, output_dir=str(output_dir))

    assert path.endswith(".pdf")
    full_text = _pdf_text(path)

    assert "INTERNAL ADDENDUM" in full_text
    assert "NOT FOR CLIENT DISTRIBUTION" in full_text
    assert "Ambiguous Sector Notice For Brief Test" in full_text
    for title in SECTION_TITLES:
        assert title in full_text, f"missing section: {title}"

    # The notice link must never be dropped from the Scouting Assessment section.
    assert "https://www.find-tender.service.gov.uk/Notice/REF-BRIEF-1" in full_text

    assert "Buyer in scope" in full_text


def test_build_brief_includes_provisional_label_when_assessment_exists(conn, tmp_path):
    notice_id = _make_notice(conn)
    save_scope_read(
        conn,
        notice_id,
        {
            "capability_fit": {"rating": "LOW", "reasoning": "No clear match."},
            "competitor_position": {"rating": "UNKNOWN", "reasoning": "Not assessed."},
            "right_to_win": {"rating": "LOW", "reasoning": "Weak fit."},
            "overall": {"rating": "DECLINE", "reasoning": "Not worth pursuing."},
            "open_questions": ["Is this really out of sector?"],
        },
    )
    path = build_brief(conn, notice_id, output_dir=str(tmp_path / "briefs"))
    full_text = _pdf_text(path)

    assert "PROVISIONAL — FOR VALIDATION" in full_text
    assert "Is this really out of sector?" in full_text


def test_build_original_notice_pdf_preserves_source_and_full_text(conn, tmp_path):
    notice_id = _make_notice(conn)
    path = build_original_notice_pdf(conn, notice_id, output_dir=str(tmp_path / "briefs"))
    full_text = _pdf_text(path)

    assert path.endswith("_original_notice.pdf")
    assert "ORIGINAL NOTICE" in full_text
    assert "https://www.find-tender.service.gov.uk/Notice/REF-BRIEF-1" in full_text
    assert "real-time payments platform integration with smart grid billing systems" in full_text


def test_build_brief_renders_risks_asks_and_recommendation(conn, tmp_path):
    notice_id = _make_notice(conn)
    save_scope_read(
        conn,
        notice_id,
        {
            "capability_fit": {"rating": "HIGH", "reasoning": "Strong engineering fit."},
            "competitor_position": {"rating": "UNKNOWN", "reasoning": "Not assessed."},
            "right_to_win": {"rating": "MED", "reasoning": "Adjacent capability."},
            "overall": {"rating": "PURSUE", "reasoning": "Closest available match."},
            "open_questions": ["Should not appear -- asks take priority."],
            "capability_mapping": [
                {"problem": "Settlement calculations", "capability_mapping": "&Money financial platform"},
            ],
            "positioning_points": [
                {
                    "point": "No UK delivery reference yet",
                    "how_to_address": "Lead with European proof points, e.g. &Money.",
                },
            ],
            "blockers": [
                {"blocker": "Framework access", "assessment": "Named call-off Trifork is not on."},
            ],
            "asks": [
                {"ask": "Confirm UK delivery capacity.", "why_it_matters": "Right to win depends on it."},
            ],
            "recommendation": {
                "decision": "PROCEED",
                "immediate_actions": ["Register interest via the buyer's portal."],
                "rationale": "Closest available match to engineering strength.",
            },
        },
    )
    path = build_brief(conn, notice_id, output_dir=str(tmp_path / "briefs"))
    full_text = _pdf_text(path)

    assert "Framework access" in full_text
    assert "Named call-off Trifork is not on." in full_text
    assert "Confirm UK delivery capacity." in full_text
    assert "PROCEED" in full_text
    assert "Should not appear" not in full_text


def test_capture_brief_has_exact_ten_sections_in_order(conn, tmp_path):
    notice_id = _make_notice(conn)
    path = build_capture_brief(conn, notice_id, output_dir=str(tmp_path / "briefs"))
    full_text = _pdf_text(path)

    positions = [full_text.index(title) for title in CAPTURE_SECTION_TITLES]
    assert positions == sorted(positions)
    assert "PROVISIONAL — FOR VALIDATION" in full_text


def test_word_internal_addendum_has_required_sections(conn, tmp_path):
    notice_id = _make_notice(conn)
    # Section C's heading now names the real rating (2026-08-19 house style
    # fix; it used to say "HIGH" unconditionally regardless of the actual
    # capability_fit), so this needs a real assessment on file to assert
    # against.
    save_scope_read(
        conn,
        notice_id,
        {
            "capability_fit": {"rating": "HIGH", "reasoning": "Strong engineering fit."},
            "competitor_position": {"rating": "UNKNOWN", "reasoning": "Not assessed."},
            "right_to_win": {"rating": "MED", "reasoning": "Adjacent capability."},
            "overall": {"rating": "PURSUE", "reasoning": "Closest available match."},
            "open_questions": [],
        },
    )
    path = build_internal_addendum_docx(conn, notice_id, str(tmp_path / "briefs"))
    full_text = _docx_text(path)

    assert path.endswith("_internal_addendum.docx")
    positions = [full_text.index(title) for title in (
        "A. TRIAGE GATE SUMMARY", "B. SCOUTING ASSESSMENT", "C. WHY THIS IS A HIGH FIT",
        "D. OPEN BLOCKERS AND RISKS", "E. DIRECT ASKS FOR TRIFORK VIA VICTORIA",
        "F. DECISION REQUESTED FROM VICTORIA",
    )]
    assert positions == sorted(positions)
    assert "INTERNAL USE ONLY" in full_text
    assert "PROVISIONAL, FOR VALIDATION" in full_text
    assert "Smarter Bids. Real Results." in full_text


def test_word_capture_brief_has_exact_ten_sections(conn, tmp_path):
    notice_id = _make_notice(conn)
    path = build_capture_brief_docx(conn, notice_id, str(tmp_path / "briefs"))
    full_text = _docx_text(path)

    assert path.endswith("_capture_brief.docx")
    headings = (
        "1. EXECUTIVE SUMMARY", "2. KEY TERMS", "3. PROCUREMENT MECHANICS",
        "4. PROCUREMENT TIMETABLE", "5. SCOPE OF REQUIREMENT",
        "6. CAPABILITY AND FIT ASSESSMENT", "7. DECISION FRAMEWORK",
        "8. IMMEDIATE ACTIONS REQUIRED", "9. SOLO OR PARTNER RECOMMENDATION",
        "10. SUMMARY DECISION PACK",
    )
    positions = [full_text.index(title) for title in headings]
    assert positions == sorted(positions)
    assert "PROVISIONAL, FOR VALIDATION" in full_text
    assert "Smarter Bids. Real Results." in full_text


def test_word_capture_brief_is_executive_and_has_live_notice_link(conn, tmp_path):
    notice_id = _make_notice(conn)
    path = build_capture_brief_docx(conn, notice_id, str(tmp_path / "briefs"))
    document = Document(path)

    banner_cell = document.tables[0].cell(0, 0)
    assert len(banner_cell.paragraphs) == 1
    assert banner_cell.paragraphs[0].paragraph_format.space_before.pt == 0
    assert "Executive view:" in document.paragraphs[4].text

    texts = [paragraph.text for paragraph in document.paragraphs]
    scope_start = texts.index("5. SCOPE OF REQUIREMENT")
    capability_start = texts.index("6. CAPABILITY AND FIT ASSESSMENT")
    scope_content = [text for text in texts[scope_start + 1:capability_start] if text]
    assert len(scope_content) >= 2
    assert all(text != "—" for text in scope_content)

    hyperlink_targets = {
        relationship.target_ref
        for relationship in document.part.rels.values()
        if relationship.reltype == "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink"
    }
    assert "https://www.find-tender.service.gov.uk/Notice/REF-BRIEF-1" in hyperlink_targets
    notice_link_cell = document.tables[4].cell(9, 1)
    assert notice_link_cell._tc.xpath(".//w:hyperlink")


def test_word_internal_addendum_asks_victoria_not_the_buyer(conn, tmp_path):
    notice_id = _make_notice(conn)
    path = build_internal_addendum_docx(conn, notice_id, str(tmp_path / "briefs"))
    document = Document(path)
    ask_text = "\n".join(cell.text for row in document.tables[6].rows[1:] for cell in row.cells)

    assert "Does Victoria approve pursuing" in ask_text
    assert "Decision required from Victoria" in ask_text
    hyperlink_targets = {
        relationship.target_ref
        for relationship in document.part.rels.values()
        if relationship.reltype == "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink"
    }
    assert "https://www.find-tender.service.gov.uk/Notice/REF-BRIEF-1" in hyperlink_targets


def test_record_brief_inserts_row(conn, tmp_path):
    notice_id = _make_notice(conn)
    path = build_brief(conn, notice_id, output_dir=str(tmp_path / "briefs"))
    brief_id = record_brief(conn, notice_id, "gate_flag:gate1", path, "system_triage")

    row = conn.execute("SELECT * FROM escalation_briefs WHERE id = ?", (brief_id,)).fetchone()
    assert row["notice_id"] == notice_id
    assert row["docx_path"] == path
    assert row["emailed_at"] is None
