import json
from datetime import date, datetime, timezone

from docx import Document

from savvy_scout.reporting.reports import (
    generate_monthly_report,
    generate_weekly_report,
    most_recent_monday,
)


def _insert_notice(conn, **overrides):
    defaults = {
        "ref": "REF-1",
        "title": "Some Opportunity",
        "buyer": "Some Council",
        "source": "Find a Tender",
        "uk_stage": "UK3",
        "status": "NEW",
        "sector": "Central and Local Government",
        "owner": "Kanvesh",
        "cpv_primary": "72212000",  # within Central and Local Government's allowed 72/48 CPV scope
        "indicative_value": "GBP 250000",
        "first_published_at": datetime.now(timezone.utc).isoformat(),
        "first_seen_at": datetime.now(timezone.utc).isoformat(),
        "last_swept_at": datetime.now(timezone.utc).isoformat(),
        "created_at": datetime.now(timezone.utc).isoformat(),
        "updated_at": datetime.now(timezone.utc).isoformat(),
        "raw_json": "{}",
    }
    defaults.update(overrides)
    cols = list(defaults.keys())
    placeholders = ", ".join("?" for _ in cols)
    cursor = conn.execute(
        f"INSERT INTO notices ({', '.join(cols)}) VALUES ({placeholders})",
        [defaults[c] for c in cols],
    )
    conn.commit()
    return cursor.lastrowid


def _all_paragraph_and_cell_text(doc: Document) -> str:
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def test_most_recent_monday_returns_a_monday():
    for offset_days in range(7):
        d = date(2026, 8, 10 + offset_days)  # 2026-08-10 is a Monday
        assert most_recent_monday(d).weekday() == 0


def test_weekly_report_includes_notice_published_in_window(conn, tmp_path):
    week_start = date(2026, 8, 10)
    _insert_notice(
        conn,
        ref="REF-IN-WINDOW",
        title="In-window opportunity",
        buyer="Some Trust",
        first_published_at="2026-08-11T09:00:00+00:00",
        status="PHASE2_SCOPED",
    )
    _insert_notice(
        conn,
        ref="REF-OUT-OF-WINDOW",
        title="Out-of-window opportunity",
        first_published_at="2026-08-01T09:00:00+00:00",
    )

    path = generate_weekly_report(conn, week_start, str(tmp_path))
    assert path.endswith("Trifork Scouting Weekly Report 2026-08-10.docx")

    text = _all_paragraph_and_cell_text(Document(path))
    assert "In-window opportunity" in text
    assert "Out-of-window opportunity" not in text
    assert "Some Trust" in text
    assert "Awaiting Phase 2 AI scope read" in text


def test_weekly_report_excludes_out_of_scope_notices(conn, tmp_path):
    # Regression: an earlier version pulled every notice published this
    # week regardless of sector match, dumping the entire raw sweep
    # (including totally unmatched buyers) into a client-facing report.
    _insert_notice(
        conn,
        ref="REF-UNMATCHED",
        title="Agricultural equipment supply",
        buyer="Agriculture and Horticulture Development Board",
        sector=None,
        owner=None,
        cpv_primary=None,
        first_published_at="2026-08-11T09:00:00+00:00",
        status="REJECTED",
    )
    path = generate_weekly_report(conn, date(2026, 8, 10), str(tmp_path))
    text = _all_paragraph_and_cell_text(Document(path))
    assert "Agriculture and Horticulture Development Board" not in text
    assert "No new opportunities were identified this week." in text


def test_weekly_report_uses_phase2_reasoning_and_product_mapping(conn, tmp_path):
    notice_id = _insert_notice(
        conn,
        ref="REF-PHASE2",
        title="Clinical data platform build",
        first_published_at="2026-08-11T09:00:00+00:00",
        status="AWAITING_PHASE2_APPROVAL",
    )
    conn.execute(
        "INSERT INTO phase2_assessments (notice_id, capability_fit_rating, capability_fit_reasoning, "
        "competitor_position_rating, competitor_position_reasoning, right_to_win_rating, "
        "right_to_win_reasoning, overall_rating, overall_reasoning, open_questions, "
        "capability_mapping, model_used, created_at) VALUES (?, 'HIGH', 'x', 'UNKNOWN', 'x', "
        "'HIGH', 'x', 'PURSUE', 'Strong clinical data fit for Corax.', '[]', ?, 'claude-sonnet-5', ?)",
        (
            notice_id,
            json.dumps([{"problem": "Clinical data integration", "capability_mapping": "Corax"}]),
            datetime.now(timezone.utc).isoformat(),
        ),
    )
    conn.commit()

    path = generate_weekly_report(conn, date(2026, 8, 10), str(tmp_path))
    text = _all_paragraph_and_cell_text(Document(path))
    assert "Strong clinical data fit for Corax." in text
    assert "Corax" in text


def test_monthly_report_handles_value_ranges_without_mashing_digits(conn, tmp_path):
    # Regression (2026-08-15): "GBP 1100000 to 1100000" was being reduced to
    # a single blob of concatenated digits ("11000001100000" -> ~£11
    # trillion) instead of parsed as two numbers and averaged.
    _insert_notice(
        conn, ref="REF-RANGE", title="Ranged value opportunity",
        indicative_value="GBP 1100000 to 1100000",
        first_published_at="2026-08-05T09:00:00+00:00",
    )
    path = generate_monthly_report(conn, date(2026, 8, 1), str(tmp_path))
    text = _all_paragraph_and_cell_text(Document(path))
    assert "£1,100,000" in text
    assert "£11,000,001,100,000" not in text
    assert "trillion" not in text.lower()


def test_monthly_report_buckets_sectors_and_decisions(conn, tmp_path):
    month_start = date(2026, 8, 1)
    _insert_notice(
        conn, ref="REF-A", title="Council platform build", sector="Central and Local Government",
        indicative_value="GBP 100000", first_published_at="2026-08-05T09:00:00+00:00",
        status="ESCALATED_TO_VICTORIA",
    )
    _insert_notice(
        conn, ref="REF-B", title="NHS data platform", sector="NHS and Healthcare",
        indicative_value="GBP 50000", first_published_at="2026-08-06T09:00:00+00:00",
        status="REJECTED",
    )

    path = generate_monthly_report(conn, month_start, str(tmp_path))
    assert path.endswith("Trifork Scouting Monthly Report 2026-08.docx")

    text = _all_paragraph_and_cell_text(Document(path))
    assert "Central and Local Government" in text
    assert "NHS and Healthcare" in text
    assert "£150,000" in text  # TOTAL row
    assert "In Progress" in text
    assert "No Bid" in text
    assert "BIDS SUBMITTED" in text
    assert "does not currently track" in text


def test_monthly_report_marks_passed_deadline_rejections_as_too_late(conn, tmp_path):
    notice_id = _insert_notice(
        conn, ref="REF-LATE", title="Missed window opportunity",
        first_published_at="2026-08-05T09:00:00+00:00", status="REJECTED",
    )
    conn.execute(
        "INSERT INTO triage_runs (notice_id, headline_gate, headline_outcome, headline_reason, evaluated_at) "
        "VALUES (?, 'gate4', 'FAIL', 'Closed or awarded (status: complete).', ?)",
        (notice_id, datetime.now(timezone.utc).isoformat()),
    )
    triage_run_id = conn.execute("SELECT last_insert_rowid() AS id").fetchone()["id"]
    conn.execute(
        "INSERT INTO gate_results (triage_run_id, notice_id, gate_number, gate_name, outcome, reason, evaluated_at) "
        "VALUES (?, ?, 'gate4', 'Window', 'FAIL', 'Closed or awarded, deadline passed.', ?)",
        (triage_run_id, notice_id, datetime.now(timezone.utc).isoformat()),
    )
    conn.commit()

    path = generate_monthly_report(conn, date(2026, 8, 1), str(tmp_path))
    text = _all_paragraph_and_cell_text(Document(path))
    assert "Too Late" in text
