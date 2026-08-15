"""Weekly New Opportunities Summary and Monthly Opportunity Pipeline Report
for Trifork's leadership team (2026-08-13, explicit request), matching the
Bid Savvy Solutions Ltd templates exactly.

Both reports read only from data the app already tracks. Two template
sections have no equivalent in the pipeline today and are always left as
guidance placeholders rather than guessed at:
- "Bids submitted" (Monthly, section 4): the state machine has no
  "submitted to buyer" milestone -- APPROVED/CAPTURE_BRIEF_DRAFTED/
  DOCS_DOWNLOADED/CALENDARED/ACTIVE are all pre-submission capture-prep
  states. Whether and when a bid was actually submitted is knowledge
  Victoria/the owner holds, not something swept or triaged.
- "Upcoming events" (Monthly, section 5) and "Any other business"
  (section 6): free-text items with no data source in this app at all.
"""

import json
import re
import sqlite3
from datetime import date, datetime, timedelta, timezone
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from savvy_scout.dashboard.scope_filter import in_scope_filter_sql

RED = RGBColor(0xAF, 0x1F, 0x23)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
DARK = RGBColor(0x11, 0x11, 0x11)
GREY = RGBColor(0x55, 0x55, 0x55)

# Named Trifork products/case studies the capability profile cites --
# matched against phase2_assessments.capability_mapping to surface a
# concrete "Product/Service Fit" line rather than a generic sector guess.
TRIFORK_PRODUCTS = ["Corax", "Tiris Messenger", "iFly4", "Trifork PIM", "LOFTHome", "Synq"]

NEXT_ACTION_BY_STATUS = {
    "NEW": "Awaiting Phase 1 triage",
    "PHASE1_TRIAGED": "Awaiting routing",
    "TO_REVIEW": "Owner to confirm Phase 1 outcome",
    "PHASE2_SCOPED": "Awaiting Phase 2 AI scope read",
    "AWAITING_PHASE2_APPROVAL": "Owner to review Phase 2 read and confirm",
    "ESCALATED_TO_VICTORIA": "Awaiting Victoria's go/no-go decision",
    "APPROVED": "Trifork to evaluate and confirm bid/no-bid",
    "CAPTURE_BRIEF_DRAFTED": "Trifork to evaluate and confirm bid/no-bid",
    "DOCS_DOWNLOADED": "Bid Savvy to register on portal / begin capture",
    "CALENDARED": "Bid preparation in progress",
    "ACTIVE": "Bid preparation in progress",
    "REJECTED": "No further action -- declined",
    "PARKED": "Parked -- awaiting more information",
    "MONITORING": "Monitoring for a future tender stage",
}

# Monthly section 2's three-way status a Trifork leadership report expects.
# MONITORING/PARKED/TO_REVIEW/PHASE2_SCOPED/AWAITING_PHASE2_APPROVAL/
# ESCALATED_TO_VICTORIA/APPROVED/CAPTURE_BRIEF_DRAFTED/DOCS_DOWNLOADED/
# CALENDARED/ACTIVE are all still live -> "In Progress". REJECTED -> "No
# Bid". "Too Late" has no direct status of its own -- it's a REJECTED
# notice whose Gate 4 reason names a passed deadline/closed window, not a
# capability/sector decision.
STATUS_TO_DECISION = {
    "REJECTED": "No Bid",
}
IN_PROGRESS_STATUSES = {
    "PHASE1_TRIAGED", "TO_REVIEW", "PHASE2_SCOPED", "AWAITING_PHASE2_APPROVAL",
    "ESCALATED_TO_VICTORIA", "APPROVED", "CAPTURE_BRIEF_DRAFTED", "DOCS_DOWNLOADED",
    "CALENDARED", "ACTIVE", "MONITORING", "PARKED",
}


def _now() -> str:
    return datetime.now(timezone.utc).isoformat()


def most_recent_monday(reference: date | None = None) -> date:
    reference = reference or datetime.now(timezone.utc).date()
    return reference - timedelta(days=reference.weekday())


def _shade_cell(cell, hex_color: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_cell_text(cell, text: str, bold: bool = False, color: RGBColor | None = None, size: int = 10) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color


def _add_masthead(doc: Document, title: str, prepared_for: str, meta_line: str, intro: str | None = None) -> None:
    p = doc.add_paragraph()
    r = p.add_run("BID SAVVY SOLUTIONS LTD")
    r.bold = True
    r.font.size = Pt(11)
    r2 = p.add_run("  |  Trifork Account")
    r2.font.size = Pt(11)
    r2.font.color.rgb = GREY

    p2 = doc.add_paragraph()
    r3 = p2.add_run("Smarter Bids. Real Results.  |  © 2026 Bid Savvy Solutions Ltd")
    r3.italic = True
    r3.font.size = Pt(9)
    r3.font.color.rgb = GREY

    h = doc.add_paragraph()
    hr = h.add_run(title)
    hr.bold = True
    hr.font.size = Pt(20)
    hr.font.color.rgb = RED

    meta = doc.add_paragraph()
    mr = meta.add_run(f"Prepared for: {prepared_for}  |  {meta_line}")
    mr.bold = True
    mr.font.size = Pt(10.5)

    if intro:
        ip = doc.add_paragraph()
        ir = ip.add_run(intro)
        ir.italic = True
        ir.font.size = Pt(9.5)
        ir.font.color.rgb = GREY

    doc.add_paragraph()


def _fetch_latest_phase2(conn: sqlite3.Connection, notice_id: int):
    return conn.execute(
        "SELECT * FROM phase2_assessments WHERE notice_id = ? ORDER BY id DESC LIMIT 1",
        (notice_id,),
    ).fetchone()


def _fetch_gate5_reason(conn: sqlite3.Connection, notice_id: int) -> str:
    row = conn.execute(
        "SELECT reason FROM gate_results WHERE notice_id = ? AND gate_number = 'gate5' "
        "ORDER BY id DESC LIMIT 1",
        (notice_id,),
    ).fetchone()
    return row["reason"] if row else ""


def _fetch_latest_gate_reason(conn: sqlite3.Connection, notice_id: int, gate_number: str) -> str:
    row = conn.execute(
        "SELECT reason FROM gate_results WHERE notice_id = ? AND gate_number = ? "
        "ORDER BY id DESC LIMIT 1",
        (notice_id, gate_number),
    ).fetchone()
    return row["reason"] if row else ""


def _product_fit(phase2_row) -> str:
    if not phase2_row or not phase2_row["capability_mapping"]:
        return "Not yet assessed -- awaiting Phase 2 scope read"
    try:
        mapping = json.loads(phase2_row["capability_mapping"])
    except (ValueError, TypeError):
        return "Not yet assessed -- awaiting Phase 2 scope read"
    matched = set()
    for entry in mapping:
        text = entry.get("capability_mapping", "")
        for product in TRIFORK_PRODUCTS:
            if product.lower() in text.lower():
                matched.add(product)
    if matched:
        return ", ".join(sorted(matched))
    return "General engineering/data capability fit (no named product match)"


def _term_text(notice_row, gate5_reason: str) -> str:
    start = notice_row["contract_start_date"]
    end = notice_row["contract_max_extent_date"]
    years_text = None
    if start and end:
        try:
            d1 = datetime.fromisoformat(start.replace("Z", "+00:00"))
            d2 = datetime.fromisoformat(end.replace("Z", "+00:00"))
            years = round((d2 - d1).days / 365.25, 1)
            if years > 0:
                years_text = f"~{years} years"
        except ValueError:
            years_text = None
    is_framework = "framework" in (gate5_reason or "").lower()
    framework_tag = "Framework" if is_framework else "Single contract"
    if years_text:
        return f"{years_text} ({framework_tag})"
    return f"Not stated in notice ({framework_tag})"


def _summary_text(notice_row, phase2_row) -> str:
    if phase2_row and phase2_row["overall_reasoning"]:
        return phase2_row["overall_reasoning"]
    sector = notice_row["sector"] or "Unclassified sector"
    buyer = notice_row["buyer"] or "Buyer not stated"
    return (
        f"{sector} opportunity from {buyer}. Not yet through a Phase 2 scope read -- "
        f"summary will be available once assessed."
    )


def generate_weekly_report(
    conn: sqlite3.Connection, week_start: date, output_dir: str, owner: str | None = None
) -> str:
    """New opportunities first identified in the 7 days starting week_start
    (Monday) -- a short-form flag only, per the template. owner, if given
    (e.g. "Mark"), restricts to that owner's sectors only (2026-08-15,
    explicit request: Mark's reports should not include Kanvesh's Central
    and Local Government notices). Filename: 'Trifork Scouting Weekly
    Report YYYY-MM-DD.docx'."""
    week_end = week_start + timedelta(days=7)
    scope_where, scope_params = in_scope_filter_sql(conn)
    owner_clause = " AND owner = ?" if owner else ""
    owner_params = [owner] if owner else []
    rows = conn.execute(
        f"SELECT * FROM notices WHERE first_published_at >= ? AND first_published_at < ? "
        f"AND ({scope_where}){owner_clause} ORDER BY sector, buyer, title",
        [week_start.isoformat(), week_end.isoformat()] + scope_params + owner_params,
    ).fetchall()

    doc = Document()
    _add_masthead(
        doc,
        "Weekly New Opportunities Summary",
        "Trifork Leadership Team",
        f"Week commencing: {week_start.strftime('%d %B %Y')}  |  Due: 16:00 each Monday",
        "This summary lists new opportunities identified during the reporting week. "
        "A full notice and briefing document (Internal Addendum) follows separately for any "
        "opportunity progressing to evaluation.",
    )

    if not rows:
        p = doc.add_paragraph()
        p.add_run("No new opportunities were identified this week.").italic = True

    for i, row in enumerate(rows, start=1):
        phase2 = _fetch_latest_phase2(conn, row["id"])
        gate5_reason = _fetch_gate5_reason(conn, row["id"])

        title_table = doc.add_table(rows=1, cols=1)
        title_table.style = "Table Grid"
        cell = title_table.rows[0].cells[0]
        _shade_cell(cell, "AF1F23")
        _set_cell_text(
            cell, f"OPPORTUNITY {i}: {row['buyer'] or 'Buyer not stated'} | {row['title']}",
            bold=True, color=WHITE, size=11,
        )

        detail = doc.add_table(rows=2, cols=2)
        detail.style = "Table Grid"
        _set_cell_text(detail.cell(0, 0), f"STAGE: {row['uk_stage']}", bold=True, size=9.5)
        _set_cell_text(
            detail.cell(0, 1),
            f"CONTRACT VALUE: {row['indicative_value'] or 'Not stated'}", bold=True, size=9.5,
        )
        _set_cell_text(detail.cell(1, 0), f"TERM: {_term_text(row, gate5_reason)}", bold=True, size=9.5)
        _set_cell_text(
            detail.cell(1, 1), f"PRODUCT / SERVICE FIT: {_product_fit(phase2)}", bold=True, size=9.5,
        )

        summary_p = doc.add_paragraph()
        summary_p.add_run("SUMMARY: ").bold = True
        summary_p.add_run(_summary_text(row, phase2))

        action_p = doc.add_paragraph()
        action_p.add_run("NEXT ACTION: ").bold = True
        action_p.add_run(NEXT_ACTION_BY_STATUS.get(row["status"], "Review required"))

        doc.add_paragraph()

    filename = f"Trifork Scouting Weekly Report {week_start.isoformat()}.docx"
    path = Path(output_dir) / filename
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    return str(path)


def generate_monthly_report(
    conn: sqlite3.Connection, month_start: date, output_dir: str, owner: str | None = None
) -> str:
    """Opportunities identified, decisions, bids in progress, and (as
    placeholders -- see module docstring) bids submitted/upcoming events/AOB,
    for the calendar month containing month_start. owner, if given, restricts
    to that owner's sectors only (see generate_weekly_report's docstring).
    Filename: 'Trifork Scouting Monthly Report YYYY-MM.docx'."""
    if month_start.month == 12:
        month_end = date(month_start.year + 1, 1, 1)
    else:
        month_end = date(month_start.year, month_start.month + 1, 1)

    scope_where, scope_params = in_scope_filter_sql(conn)
    owner_clause = " AND owner = ?" if owner else ""
    owner_params = [owner] if owner else []
    rows = conn.execute(
        f"SELECT * FROM notices WHERE first_published_at >= ? AND first_published_at < ? "
        f"AND ({scope_where}){owner_clause} ORDER BY sector, buyer, title",
        [month_start.isoformat(), month_end.isoformat()] + scope_params + owner_params,
    ).fetchall()

    doc = Document()
    _add_masthead(
        doc,
        "Monthly Opportunity Pipeline Report",
        "Trifork Leadership Team",
        f"Reporting period: {month_start.strftime('%B %Y')}",
        "Prepared under the Bid Savvy Solutions Ltd monthly retainer, covering opportunity "
        "identification and capture planning for the period stated above.",
    )

    # --- Section 1: opportunities identified this month, by sector -----
    h1 = doc.add_paragraph()
    h1.add_run("1. OPPORTUNITIES IDENTIFIED THIS MONTH").bold = True
    by_sector: dict[str, list] = {}
    for row in rows:
        by_sector.setdefault(row["sector"] or "Unclassified", []).append(row)

    def _value_to_float(value_text):
        # Regression (2026-08-15): indicative_value isn't always a single
        # number -- OCDS ranges look like "GBP 1100000 to 1100000". Blindly
        # concatenating every digit character (the old approach) mashed
        # both numbers of a range into one, producing nonsense totals in
        # the hundreds of billions. Extract each separate numeric token and
        # average a two-number range instead.
        if not value_text:
            return 0.0
        tokens = re.findall(r"\d+\.?\d*", value_text)
        if not tokens:
            return 0.0
        numbers = [float(t) for t in tokens]
        return sum(numbers) / len(numbers)

    t1 = doc.add_table(rows=1, cols=3)
    t1.style = "Table Grid"
    for idx, header in enumerate(["SECTOR", "NO. OF OPPORTUNITIES", "TOTAL VALUE"]):
        _shade_cell(t1.rows[0].cells[idx], "AF1F23")
        _set_cell_text(t1.rows[0].cells[idx], header, bold=True, color=WHITE)
    total_n, total_v = 0, 0.0
    for sector, sector_rows in sorted(by_sector.items()):
        sector_value = sum(_value_to_float(r["indicative_value"]) for r in sector_rows)
        total_n += len(sector_rows)
        total_v += sector_value
        r = t1.add_row().cells
        _set_cell_text(r[0], sector)
        _set_cell_text(r[1], str(len(sector_rows)))
        _set_cell_text(r[2], f"£{sector_value:,.0f}" if sector_value else "Not stated")
    r = t1.add_row().cells
    _set_cell_text(r[0], "TOTAL", bold=True)
    _set_cell_text(r[1], str(total_n), bold=True)
    _set_cell_text(r[2], f"£{total_v:,.0f}" if total_v else "Not stated", bold=True)
    doc.add_paragraph()

    # --- Section 2: decisions and status --------------------------------
    h2 = doc.add_paragraph()
    h2.add_run("2. DECISIONS AND STATUS").bold = True
    t2 = doc.add_table(rows=1, cols=4)
    t2.style = "Table Grid"
    for idx, header in enumerate(["CLIENT", "OPPORTUNITY TITLE", "STATUS", "RATIONALE"]):
        _shade_cell(t2.rows[0].cells[idx], "AF1F23")
        _set_cell_text(t2.rows[0].cells[idx], header, bold=True, color=WHITE)
    for row in rows:
        if row["status"] == "REJECTED":
            gate_reason = (
                _fetch_latest_gate_reason(conn, row["id"], "gate1")
                or _fetch_latest_gate_reason(conn, row["id"], "gate2")
                or "Declined at Phase 1/2 review"
            )
            gate4_reason = _fetch_latest_gate_reason(conn, row["id"], "gate4")
            is_too_late = "passed" in (gate4_reason or "").lower() or "closed" in (gate4_reason or "").lower()
            status_label = "Too Late" if is_too_late else "No Bid"
            rationale = gate4_reason if is_too_late else gate_reason
        elif row["status"] in IN_PROGRESS_STATUSES:
            status_label = "In Progress"
            rationale = NEXT_ACTION_BY_STATUS.get(row["status"], "")
        else:
            continue
        r = t2.add_row().cells
        _set_cell_text(r[0], row["buyer"] or "Not stated")
        _set_cell_text(r[1], row["title"])
        _set_cell_text(r[2], status_label, bold=True)
        _set_cell_text(r[3], (rationale or "")[:200])
    doc.add_paragraph()

    # --- Section 3: bids in progress -------------------------------------
    h3 = doc.add_paragraph()
    h3.add_run("3. BIDS IN PROGRESS").bold = True
    t3 = doc.add_table(rows=1, cols=4)
    t3.style = "Table Grid"
    for idx, header in enumerate(["CLIENT", "OPPORTUNITY TITLE", "STAGE", "VALUE"]):
        _shade_cell(t3.rows[0].cells[idx], "AF1F23")
        _set_cell_text(t3.rows[0].cells[idx], header, bold=True, color=WHITE)
    in_progress_rows = [
        r for r in rows
        if r["status"] in {
            "APPROVED", "CAPTURE_BRIEF_DRAFTED", "DOCS_DOWNLOADED", "CALENDARED", "ACTIVE",
            "ESCALATED_TO_VICTORIA", "AWAITING_PHASE2_APPROVAL",
        }
    ]
    for row in in_progress_rows:
        r = t3.add_row().cells
        _set_cell_text(r[0], row["buyer"] or "Not stated")
        _set_cell_text(r[1], row["title"])
        _set_cell_text(r[2], row["uk_stage"])
        _set_cell_text(r[3], row["indicative_value"] or "Not stated")
    if not in_progress_rows:
        r = t3.add_row().cells
        _set_cell_text(r[0], "None this period")
        for c in r[1:]:
            _set_cell_text(c, "")
    doc.add_paragraph()

    # --- Section 4: bids submitted (placeholder -- not tracked) ---------
    h4 = doc.add_paragraph()
    h4.add_run("4. BIDS SUBMITTED").bold = True
    p4 = doc.add_paragraph()
    p4.add_run(
        "Savvy Scout does not currently track a \"submitted to buyer\" milestone -- "
        "please complete this section manually with any bids submitted this period."
    ).italic = True
    doc.add_paragraph()

    # --- Section 5: upcoming events (placeholder) ------------------------
    h5 = doc.add_paragraph()
    h5.add_run("5. UPCOMING EVENTS").bold = True
    p5 = doc.add_paragraph()
    p5.add_run("[Event Name] | [DD Month] | [Why this matters to Trifork]").italic = True
    doc.add_paragraph()

    # --- Section 6: any other business (placeholder) ---------------------
    h6 = doc.add_paragraph()
    h6.add_run("6. ANY OTHER BUSINESS").bold = True
    p6 = doc.add_paragraph()
    p6.add_run(
        "Matters for Trifork's attention that fall outside the sections above, for example "
        "resourcing, risks, or strategic considerations."
    )
    b1 = doc.add_paragraph(style="List Bullet")
    b1.add_run("[First matter for attention.]")
    b2 = doc.add_paragraph(style="List Bullet")
    b2.add_run("[Second matter for attention, if applicable.]")

    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.add_run("Prepared by: Victoria Milan, Bid Director, Bid Savvy Solutions Ltd").bold = True

    filename = f"Trifork Scouting Monthly Report {month_start.strftime('%Y-%m')}.docx"
    path = Path(output_dir) / filename
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    return str(path)
