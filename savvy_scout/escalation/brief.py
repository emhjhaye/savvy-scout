"""Escalation documents (SPEC v1.5 B3).

- Internal Addendum: generated on owner-marked Victoria escalation.
- Capture Brief: generated only after Victoria GO (approve).
"""

import json
import os
import sqlite3
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.shared import Pt

GATE_ORDER = ["gate1", "gate2", "gate3", "gate4", "gate5", "filter3"]

# SAVVY_SCOUT_BRIEFS_DIR lets a production deploy (Render's persistent disk,
# e.g. /var/data/briefs) point generated .docx files somewhere that actually
# survives a redeploy -- the default keeps local dev unchanged.
DEFAULT_BRIEFS_DIR = os.environ.get("SAVVY_SCOUT_BRIEFS_DIR", "briefs")


def _now() -> str:
    return datetime.now(timezone.utc).isoformat()


def _latest_gate_results(conn: sqlite3.Connection, notice_id: int) -> list[sqlite3.Row]:
    run = conn.execute(
        "SELECT id FROM triage_runs WHERE notice_id = ? ORDER BY id DESC LIMIT 1", (notice_id,)
    ).fetchone()
    if not run:
        return []
    rows = conn.execute(
        "SELECT gate_number, gate_name, outcome, reason FROM gate_results WHERE triage_run_id = ?",
        (run["id"],),
    ).fetchall()
    by_gate = {r["gate_number"]: r for r in rows}
    return [by_gate[g] for g in GATE_ORDER if g in by_gate]


def _latest_assessment(conn: sqlite3.Connection, notice_id: int) -> sqlite3.Row | None:
    return conn.execute(
        "SELECT * FROM phase2_assessments WHERE notice_id = ? ORDER BY id DESC LIMIT 1",
        (notice_id,),
    ).fetchone()


def _add_section(doc: Document, number: int, title: str, body_lines: list[str]) -> None:
    doc.add_heading(f"{number}. {title}", level=1)
    for line in body_lines:
        if line:
            doc.add_paragraph(line)


def build_internal_addendum(conn: sqlite3.Connection, notice_id: int, output_dir: str = DEFAULT_BRIEFS_DIR) -> str:
    notice = conn.execute("SELECT * FROM notices WHERE id = ?", (notice_id,)).fetchone()
    if notice is None:
        raise ValueError(f"No notice with id {notice_id}")

    gate_rows = _latest_gate_results(conn, notice_id)
    assessment = _latest_assessment(conn, notice_id)
    flagged_gates = [g for g in gate_rows if g["outcome"] in ("FLAG", "MAYBE")]

    doc = Document()
    doc.add_heading(f"INTERNAL ADDENDUM: {notice['title']}", level=0)
    warning = doc.add_paragraph()
    warning_run = warning.add_run(
        "AUTO-GENERATED PROVISIONAL DRAFT FOR VALIDATION. Internal document only, "
        "not for client distribution. Decision belongs to Victoria Milan, Bid Director."
    )
    warning_run.bold = True
    warning_run.font.size = Pt(11)

    _add_section(
        doc, 1, "Opportunity summary",
        [notice["title"], f"Reference: {notice['ref']}", f"Source: {notice['source']}"],
    )
    _add_section(
        doc, 2, "Buyer",
        [
            notice["buyer"] or "UNVERIFIED",
            notice["buyer_address"] or "",
            f"PPON: {notice['buyer_ppon']}" if notice["buyer_ppon"] else "",
            f"Organisation type: {notice['buyer_org_type']}" if notice["buyer_org_type"] else "",
        ],
    )
    _add_section(
        doc, 3, "Value",
        [
            notice["indicative_value"] or "UNVERIFIED",
            f"Inc. VAT: {notice['value_amount_gross']}" if notice["value_amount_gross"] else "",
            f"Above threshold: {'Yes' if notice['above_threshold'] else 'No'}"
            if notice["above_threshold"] is not None else "",
        ],
    )

    framework_gate = next((g for g in gate_rows if g["gate_number"] == "gate3"), None)
    _add_section(
        doc, 4, "Route to market",
        [f"UK stage: {notice['uk_stage']}", framework_gate["reason"] if framework_gate else "UNVERIFIED"],
    )

    _add_section(
        doc, 5, "Award criteria",
        [notice["award_criteria_summary"]] if notice["award_criteria_summary"] else ["Not stated."],
    )

    _add_section(
        doc, 6, "Contract dates and extensions",
        [
            f"{notice['contract_start_date'] or 'UNVERIFIED'} to {notice['contract_end_date'] or 'UNVERIFIED'}",
            f"Max extent: {notice['contract_max_extent_date']}" if notice["contract_max_extent_date"] else "",
            notice["renewal_description"] or "",
        ],
    )

    _add_section(
        doc, 7, "Gate outcomes",
        [f"{g['gate_name']}: {g['outcome']} -- {g['reason']}" for g in gate_rows]
        or ["Not yet triaged."],
    )

    if assessment:
        ratings_lines = [
            f"Capability fit: {assessment['capability_fit_rating']} -- "
            f"{assessment['capability_fit_reasoning']} (PROVISIONAL, FOR VALIDATION)",
            f"Right to win: {assessment['right_to_win_rating']} -- "
            f"{assessment['right_to_win_reasoning']} (PROVISIONAL, FOR VALIDATION)",
            f"Overall: {assessment['overall_rating']} -- "
            f"{assessment['overall_reasoning']} (PROVISIONAL, FOR VALIDATION)",
        ]
        competitor_lines = [
            f"{assessment['competitor_position_rating']} -- {assessment['competitor_position_reasoning']} "
            "(PROVISIONAL, FOR VALIDATION)"
        ]
        open_questions = json.loads(assessment["open_questions"])
    else:
        ratings_lines = ["Phase 2 scope read has not run yet for this notice."]
        competitor_lines = ["Not yet assessed."]
        open_questions = []

    _add_section(doc, 8, "Provisional ratings with reasoning", ratings_lines)
    _add_section(doc, 9, "Competitor picture", competitor_lines)
    _add_section(
        doc, 10, "Risks",
        [f"{g['gate_name']}: {g['reason']}" for g in flagged_gates] or ["No gate flags recorded."],
    )
    _add_section(
        doc, 11, "Open questions",
        open_questions + [f"{g['gate_name']}: {g['reason']}" for g in flagged_gates]
        or ["None recorded."],
    )
    _add_section(doc, 12, "Decision requested", ["Victoria decision required: go, no-go or park."])

    Path(output_dir).mkdir(parents=True, exist_ok=True)
    safe_ref = notice["ref"].replace("/", "-")
    output_path = str(Path(output_dir) / f"{safe_ref}_internal_addendum.docx")
    doc.save(output_path)
    return output_path


def build_capture_brief(conn: sqlite3.Connection, notice_id: int, output_dir: str = DEFAULT_BRIEFS_DIR) -> str:
    notice = conn.execute("SELECT * FROM notices WHERE id = ?", (notice_id,)).fetchone()
    if notice is None:
        raise ValueError(f"No notice with id {notice_id}")

    assessment = _latest_assessment(conn, notice_id)
    doc = Document()
    doc.add_heading(f"CAPTURE BRIEF: {notice['title']}", level=0)
    p = doc.add_paragraph("Prepared under Victoria Milan, Bid Director. Client-facing draft.")
    p.runs[0].bold = True

    sections = [
        ("Executive Summary", [notice["title"]]),
        ("Key Terms", [
            f"Reference: {notice['ref']}",
            f"Buyer: {notice['buyer'] or 'UNVERIFIED'}",
            f"Value: {notice['indicative_value'] or 'UNVERIFIED'}"
            + (f" (inc. VAT: {notice['value_amount_gross']})" if notice["value_amount_gross"] else ""),
            f"Above threshold: {'Yes' if notice['above_threshold'] else 'No'}"
            if notice["above_threshold"] is not None else "",
            f"Category: {notice['main_procurement_category']}" if notice["main_procurement_category"] else "",
        ]),
        ("Procurement Mechanics", [
            f"Source: {notice['source']}",
            f"UK stage: {notice['uk_stage']}",
            f"Procedure: {notice['procurement_method'] or 'UNVERIFIED'}",
            f"Procedure details: {notice['procurement_method_details']}"
            if notice["procurement_method_details"] else "",
            f"Procedure features: {notice['procedure_features']}" if notice["procedure_features"] else "",
        ]),
        ("Procurement Timetable", [
            f"Enquiry period end: {notice['enquiry_period_end']}" if notice["enquiry_period_end"] else "",
            f"Deadline: {notice['deadline'] or 'UNVERIFIED'}",
            f"Award decision (est.): {notice['award_period_end']}" if notice["award_period_end"] else "",
            f"Contract: {notice['contract_start_date'] or 'UNVERIFIED'} to "
            f"{notice['contract_end_date'] or 'UNVERIFIED'}"
            + (f" (max extent {notice['contract_max_extent_date']})" if notice["contract_max_extent_date"] else ""),
            notice["renewal_description"] or "",
        ]),
        ("Award Criteria", [notice["award_criteria_summary"] or "Not stated."]),
        ("Submission Details", [
            notice["submission_method_details"] or "Not stated -- see original notice.",
            f"Electronic submission: {notice['electronic_submission_policy']}"
            if notice["electronic_submission_policy"] else "",
            f"Languages: {notice['submission_languages']}" if notice["submission_languages"] else "",
        ]),
        ("Scope of Requirement", [notice["text_blob"][:1200] if notice["text_blob"] else "UNVERIFIED"]),
        ("Contacts and Buyer Details", [
            f"Contracting authority: {notice['buyer_address'] or 'UNVERIFIED'}",
            f"PPON: {notice['buyer_ppon']}" if notice["buyer_ppon"] else "",
            f"Contact: {notice['buyer_contact_email']}" if notice["buyer_contact_email"] else "",
            f"Website: {notice['buyer_website']}" if notice["buyer_website"] else "",
            f"Organisation type: {notice['buyer_org_type']}" if notice["buyer_org_type"] else "",
            f"Incumbent supplier: {notice['supplier_name']}" if notice["supplier_name"] else "",
            f"Conflicts assessment: {notice['conflicts_assessment']}" if notice["conflicts_assessment"] else "",
        ]),
        ("Capability and Fit Assessment", [
            f"Overall: {assessment['overall_rating']}" if assessment else "PROVISIONAL, FOR VALIDATION",
            assessment['overall_reasoning'] if assessment else "No assessment available.",
        ]),
        ("Decision Framework", ["Proceed based on Victoria GO decision and owner readiness."]),
        ("Immediate Actions Required", ["Confirm route to market, timeline and clarifications."]),
        ("Solo or Partner Recommendation", ["To be validated by Victoria and owner."]),
        ("Summary Decision Pack", ["Generated post-GO per v1.5 policy."]),
    ]
    for idx, (title, lines) in enumerate(sections, start=1):
        _add_section(doc, idx, title, lines)

    Path(output_dir).mkdir(parents=True, exist_ok=True)
    safe_ref = notice["ref"].replace("/", "-")
    output_path = str(Path(output_dir) / f"{safe_ref}_capture_brief.docx")
    doc.save(output_path)
    return output_path


def build_brief(conn: sqlite3.Connection, notice_id: int, output_dir: str = DEFAULT_BRIEFS_DIR) -> str:
    # Backward-compatible alias used by older callers.
    return build_internal_addendum(conn, notice_id, output_dir)


def record_brief(
    conn: sqlite3.Connection,
    notice_id: int,
    trigger_reason: str,
    docx_path: str,
    created_by: str,
    brief_type: str = "INTERNAL_ADDENDUM",
) -> int:
    now = _now()
    cursor = conn.execute(
        "INSERT INTO escalation_briefs (notice_id, trigger_reason, docx_path, created_by, created_at, brief_type) "
        "VALUES (?, ?, ?, ?, ?, ?)",
        (notice_id, trigger_reason, docx_path, created_by, now, brief_type),
    )
    conn.commit()
    return cursor.lastrowid


def mark_emailed(conn: sqlite3.Connection, brief_id: int, recipient: str) -> None:
    conn.execute(
        "UPDATE escalation_briefs SET emailed_to = ?, emailed_at = ? WHERE id = ?",
        (recipient, _now(), brief_id),
    )
    conn.commit()
