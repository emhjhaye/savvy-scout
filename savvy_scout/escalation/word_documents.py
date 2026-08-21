import re
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.shared import Pt
from docx.text.paragraph import Paragraph

from savvy_scout.escalation.context import MISSING, build_context

TEMPLATE_DIR = Path(__file__).resolve().parent.parent / "templates" / "artifacts"

# House style (2026-08-19 spec): exact hex values for rating text and
# callout fills/borders.
_RATING_COLORS = {
    "HIGH": "2F5D8A",
    "MED": "3A3A3A",
    "MEDIUM": "3A3A3A",
    "LOW": "AF0000",
    "UNKNOWN": "3A3A3A",
    "N/A": "3A3A3A",
}
_CALLOUT_BLUE = {"fill": "EEF3FA", "border": "2F5D8A"}
_CALLOUT_RED = {"fill": "FDF0F0", "border": "AF0000"}


def _rating_color(rating):
    return _RATING_COLORS.get(str(rating).upper(), "3A3A3A")


def _color_rating_run(paragraph, rating):
    """Colours the LAST run of a paragraph per house style's rating colour
    table, for cells whose whole text is a rating (or ends with one, e.g.
    "Capability fit: MED")."""
    if not paragraph.runs:
        return
    run = paragraph.runs[-1]
    run.font.color.rgb = None
    from docx.shared import RGBColor

    run.font.color.rgb = RGBColor.from_string(_rating_color(rating))
    run.font.bold = True


def _shade_cell(cell, fill, border):
    """Applies a callout fill + single-line border on all four sides,
    matching the house style spec's blue ("why this matters") and red
    (deadline/confidentiality) callouts."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)
    borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        edge = OxmlElement(f"w:{side}")
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), "4")
        edge.set(qn("w:color"), border)
        borders.append(edge)
    tc_pr.append(borders)

# House style (2026-08-19 spec): only these four verified case studies exist.
# Names used in error before and corrected -- must never appear again even
# if the AI slips, so this is enforced here as a defensive check, not just
# an instruction to the model.
_BANNED_CASE_STUDIES = ("vocalink", "visa", "danske bank")


def _strip_banned_names(text):
    if not isinstance(text, str):
        return text
    lowered = text.casefold()
    if any(banned in lowered for banned in _BANNED_CASE_STUDIES):
        return "UNVERIFIED, the model referenced an unverified case study, removed here."
    return text


def _house_style(text):
    """Never-invented content still needs to honour house style even if the
    model slips: no em/en dash. Skips the bare MISSING sentinel itself
    ("—"), which is an em dash used throughout the app as the "no data"
    placeholder, not house-style prose."""
    if not isinstance(text, str) or text == MISSING:
        return text
    text = _strip_banned_names(text)
    return text.replace("—", ",").replace("–", ",")


# House style (2026-08-19 spec) names Inter as the document font. python-docx
# can only write the font NAME into the file; it cannot embed the font or
# guarantee it renders that way. Whether it actually displays as Inter
# depends entirely on Inter being installed on whichever computer opens the
# .docx in Word -- if it isn't, Word silently substitutes its own fallback
# and there is no way to control that from the generating side.
_HOUSE_FONT = "Inter"


def _apply_font(paragraph):
    for run in paragraph.runs:
        run.font.name = _HOUSE_FONT
        r_pr = run._element.get_or_add_rPr()
        r_fonts = r_pr.find(qn("w:rFonts"))
        if r_fonts is None:
            r_fonts = OxmlElement("w:rFonts")
            r_pr.append(r_fonts)
        for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
            r_fonts.set(qn(attr), _HOUSE_FONT)


def _set_cell(cell, value):
    text = _house_style(str(value if value not in (None, "") else MISSING))
    paragraph = cell.paragraphs[0]
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)
    for extra in cell.paragraphs[1:]:
        extra._element.getparent().remove(extra._element)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1
    _apply_font(paragraph)


def _set_paragraph(paragraph, value):
    text = _house_style(str(value if value not in (None, "") else MISSING))
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)
    _apply_font(paragraph)


def _remove_paragraph(paragraph):
    paragraph._element.getparent().remove(paragraph._element)


def _insert_paragraph_before(document, anchor_element, text, bold=False, bullet=False):
    """Inserts a new paragraph directly before anchor_element (a table's or
    paragraph's own oxml element) -- used for content the fixed templates
    have no placeholder slot for, e.g. the MED dual-reading structure and
    Section D's mandatory verbatim framing text."""
    new_p = OxmlElement("w:p")
    anchor_element.addprevious(new_p)
    paragraph = Paragraph(new_p, document)
    run = paragraph.add_run(("• " if bullet else "") + _house_style(text))
    run.bold = bold
    _apply_font(paragraph)
    return paragraph


def _ensure_table_rows(table, needed_rows):
    """Adds rows to a fixed-shape template table when the house-style
    spec needs one more field than the template has (e.g. Contact),
    copying the last existing row's cell formatting (shading, borders) so
    the new row doesn't render as plain, unstyled cells."""
    from copy import deepcopy

    while len(table.rows) < needed_rows:
        template_row = table.rows[-1]._tr
        new_row = deepcopy(template_row)
        for cell in new_row.findall(qn("w:tc")):
            for text_el in cell.findall(f".//{qn('w:t')}"):
                text_el.text = ""
        table._tbl.append(new_row)


def _insert_med_dual_reading(document, anchor_element, context, include_reasoning=True):
    """House style (2026-08-19): a MED capability_fit is genuinely ambiguous
    between a build and a packaged-product purchase and must present both
    readings as their own structure, not one averaged reasoning sentence.
    Inserted directly before anchor_element (a table's or paragraph's own
    oxml element) since neither template has a placeholder for this --
    only rendered when capability_fit is actually MED and the model
    produced real dual-reading content."""
    if context["ai_read"]["capability_fit"] != "MED":
        return
    dual = context["med_dual_reading"]
    build_signals = dual.get("build_signals") or []
    product_signals = dual.get("product_signals") or []
    honest_position = dual.get("honest_position")
    if not (build_signals or product_signals or honest_position):
        return
    anchor = anchor_element
    reasoning = include_reasoning and context["ai_read"].get("per_field_reasoning", {}).get("capability_fit")
    if reasoning and reasoning != MISSING:
        _insert_paragraph_before(document, anchor, reasoning)
    if build_signals:
        _insert_paragraph_before(document, anchor, "Signals that point to a build:", bold=True)
        for signal in build_signals:
            _insert_paragraph_before(document, anchor, signal, bullet=True)
    if product_signals:
        _insert_paragraph_before(document, anchor, "Signals that point to a packaged product purchase:", bold=True)
        for signal in product_signals:
            _insert_paragraph_before(document, anchor, signal, bullet=True)
    if honest_position:
        _insert_paragraph_before(document, anchor, f"The honest position: {honest_position}")


def _set_hyperlink(cell, label, url):
    _set_cell(cell, "")
    paragraph = cell.paragraphs[0]
    for child in list(paragraph._element):
        paragraph._element.remove(child)
    relationship_id = paragraph.part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    colour = OxmlElement("w:color")
    colour.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    fonts = OxmlElement("w:rFonts")
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        fonts.set(qn(attr), _HOUSE_FONT)
    properties.extend((fonts, colour, underline))
    text = OxmlElement("w:t")
    text.text = label
    run.extend((properties, text))
    hyperlink.append(run)
    paragraph._element.append(hyperlink)


def _item_text(value, keys):
    if isinstance(value, dict):
        parts = [str(value[key]) for key in keys if value.get(key)]
        return " — ".join(parts) or MISSING
    return str(value or MISSING)


def _split_ask(item):
    if isinstance(item, dict):
        return item.get("ask", MISSING), item.get("why_it_matters", MISSING)
    return str(item), "Decision required from Victoria"


def _split_blocker(item):
    if isinstance(item, dict):
        return item.get("blocker", MISSING), item.get("assessment", MISSING)
    return str(item), MISSING


def _trim_table_rows(table, last_row_to_keep):
    """Removes trailing template rows beyond last_row_to_keep (0-indexed,
    inclusive) so a short real list doesn't leave a wall of "--" placeholder
    rows -- e.g. a genuinely empty blockers list should read as one clear
    statement, not six blank "Risk N" rows."""
    for row in list(table.rows[last_row_to_keep + 1:]):
        table._tbl.remove(row._tr)


def _sentence_case(value):
    text = " ".join(str(value or "").split())
    if not text:
        return ""
    text = text[0].upper() + text[1:]
    replacements = {"orr": "ORR", "uk": "UK", "it": "IT", "ai": "AI", "nhs": "NHS", "microsoft": "Microsoft", "trifork": "Trifork"}
    for source, replacement in replacements.items():
        text = re.sub(rf"\b{source}\b", replacement, text, flags=re.IGNORECASE)
    return text


def _first_sentences(value, count=1, limit=360):
    text = " ".join(str(value or MISSING).split())
    sentences = re.split(r"(?<=[.!?])\s+", text)
    result = " ".join(sentences[:count])
    if len(result) <= limit:
        return result
    return result[: limit - 1].rsplit(" ", 1)[0].rstrip(" ,;:-") + "…"


def _scope_points(context):
    title = " ".join(context["title"].casefold().split())
    chunks = []
    for line in str(context["notice_text"]).splitlines():
        cleaned = " ".join(line.split()).strip(" .")
        if not cleaned or cleaned.casefold() == title:
            continue
        if any(cleaned.casefold().startswith(prefix) for prefix in ("below threshold", "open procedure", "it services:")):
            continue
        chunks.extend(re.split(r"(?<=[.!?])\s+|;\s*", cleaned))
    points = []
    for chunk in chunks:
        polished = _sentence_case(chunk.strip(" ."))
        if len(polished) < 25 or polished.casefold() in {point.casefold() for point in points}:
            continue
        points.append(polished + ("" if polished.endswith((".", "?", "!")) else "."))
        if len(points) == 7:
            break
    return points or ["Detailed scope is not stated in the published notice."]


def _requirement_areas(context):
    areas = context["scope_of_requirement"].get("requirement_areas")
    if areas:
        return areas
    return _scope_points(context)


def _what_buyer_seeking(context):
    text = context["scope_of_requirement"].get("what_buyer_is_seeking")
    if text:
        return text
    return f"{context['route_to_market']} route to market." if context["route_to_market"] != MISSING else "Not stated in the notice."


def _engagement_model_text(context):
    model = context["engagement_model"]
    if model.get("model") or model.get("how_to_respond"):
        model_text = model.get("model") or MISSING
        how_to = model.get("how_to_respond") or MISSING
        return f"{model_text}. {how_to}"
    return "UNVERIFIED — the notice does not state how a bidder responds to this engagement."


def _key_terms_rows(context):
    return [(item.get("term", MISSING), item.get("meaning", MISSING)) for item in context["key_terms"]]


def _decision_framework_rows(context):
    rows = context["decision_framework"]
    if rows:
        return [(row.get("question", MISSING), row.get("implication", MISSING)) for row in rows]
    return [_split_ask(ask) for ask in _victoria_asks_or_ai(context)]


def _victoria_asks_or_ai(context):
    return context["direct_asks"] or _victoria_asks(context)


def _timetable_rows(context):
    ai_timetable = context["procurement_timetable_ai"]
    if ai_timetable:
        return [(row.get("milestone", MISSING), row.get("date", MISSING)) for row in ai_timetable]
    return [
        ("Notice published", context["published_date"]),
        ("Clarification deadline", context["clarification_deadline"]),
        ("Submission deadline", context["submission_deadline"]),
    ]


def _action_text(item):
    # Legacy assessments (generated before 2026-08-16) stored immediate_actions
    # as plain strings with no owner/deadline; a re-run Phase 2 read produces
    # {action, owner_and_deadline} objects instead. Handle both so an
    # un-rescoped notice's document doesn't crash.
    if isinstance(item, dict):
        return item.get("action", MISSING), item.get("owner_and_deadline", MISSING)
    return str(item), MISSING


def _immediate_action_rows(context):
    actions = context["immediate_actions"]
    if actions:
        return [_action_text(item) for item in actions]
    return [(_recommended_action(context), f"{context['owner_name']}, Bid Savvy Solutions Ltd")]


def _capability_mapping_rows(context):
    mapping = context["capability_mapping"]
    if mapping:
        return [(row.get("problem", MISSING), row.get("capability_mapping", MISSING)) for row in mapping]
    reasoning = context["ai_read"].get("per_field_reasoning", {})
    return [(
        "Capability fit assessment",
        f"{context['ai_read']['capability_fit']} — {reasoning.get('capability_fit', MISSING)}",
    )]


def _pipeline_reference(context):
    return f"N{context['notice_id']:03d}"


def _tracker_status_text(context):
    overall = context["ai_read"]["overall"]
    sheet = "Pass" if overall == "PURSUE" else "Flag"
    return (
        f"{overall}, {context['ai_read']['capability_fit']} capability fit, "
        f"in {sheet} tab, My_Trifork_Pipeline_Tracker.xlsx"
    )


def _client_brief_status_text(context):
    return (
        f"Drafted {context['generated_at'][:10]}. PROVISIONAL, awaiting Victoria's "
        f"GO / NO-GO / Park decision before any release to Trifork."
    )


def _notice_reference_with_source(context):
    if context["source_portal"] != MISSING:
        return f"{context['notice_reference']}, {context['source_portal']}"
    return context["notice_reference"]


def _final_decision_text(context):
    actions = context["immediate_actions"]
    action_text = ""
    if actions:
        numbered = " ".join(
            f"({i}) {_action_text(item)[0].rstrip('.')}." for i, item in enumerate(actions, start=1)
        )
        action_text = f" If GO, immediate actions are: {numbered}"
    rationale = context["recommendation_rationale"]
    reasoning = f" {rationale}" if rationale != MISSING else ""
    return (
        f"Decision required: GO, NO-GO or Park for {context['title']} "
        f"(ref {context['notice_reference']}).{action_text} Owner recommendation: "
        f"{_recommended_action(context)}{reasoning} — {context['owner_name']}."
    )


_NOTICE_TYPE_LABELS = {
    "UK1": "UK1 Pipeline notice",
    "UK2": "UK2 Preliminary Market Engagement",
    "UK3": "UK3 Planned procurement notice",
    "UK4": "UK4 Tender notice",
    "UK5": "UK5 Award notice",
}


def _notice_type_label(context):
    label = _NOTICE_TYPE_LABELS.get(context["uk_stage"])
    if label:
        return label
    return context["notice_type"] if context["notice_type"] != MISSING else "Not stated in the notice"


def _route_to_market_text(context):
    if context["route_to_market"] != MISSING:
        return context["route_to_market"]
    return context["engagement_model"].get("model") or MISSING


def _sources_text(context):
    parts = [context["source_portal"]]
    if context["notice_reference"] != MISSING:
        parts.append(f"reference {context['notice_reference']}")
    if context["published_date"] != MISSING:
        parts.append(f"published {context['published_date']}")
    return ", ".join(part for part in parts if part and part != MISSING)


def _urgency_text(context):
    # context["urgency"] leads with an emoji marker for dashboard scanning
    # (see derive_urgency) -- these documents are plain professional prose,
    # matching the sample, which uses no emoji anywhere.
    return context["urgency"].split(" ", 1)[-1]


def _format_value(context):
    if context["value_estimate"] == MISSING or str(context["value_estimate"]) in ("0", "0.0"):
        return "Not stated"
    try:
        amount = float(context["value_estimate"])
        formatted = f"{amount:,.0f}" if amount.is_integer() else f"{amount:,.2f}"
    except (TypeError, ValueError):
        return str(context["value_estimate"])
    symbol = {"GBP": "£", "EUR": "€", "USD": "$"}.get(context["currency"])
    return f"{symbol}{formatted}" if symbol else f"{formatted} {context['currency']}"


def _solo_or_partner_text(context):
    rec = context["solo_or_partner_recommendation"]
    if rec.get("recommendation") or rec.get("rationale"):
        return rec.get("recommendation") or MISSING, rec.get("rationale") or MISSING
    return (
        "Not yet assessed.",
        "Revisit once Trifork's UK delivery capacity and any framework/partnering "
        "requirements for this opportunity are confirmed.",
    )


def _recommended_action(context):
    overall = context["ai_read"]["overall"]
    if overall == "PURSUE":
        return "Recommend GO to capture, subject to Victoria's approval and confirmation of the open risks."
    if overall == "DECLINE":
        return "Recommend NO-GO unless Victoria accepts the identified capability and right-to-win gaps."
    return "Keep at FLAG pending Victoria's GO / NO-GO / Park decision on the identified capability and right-to-win gaps."


def _victoria_asks(context):
    rating = context["ai_read"]
    asks = [
        f"Does Victoria approve pursuing this opportunity with {rating['capability_fit']} capability fit and {rating['right_to_win']} right to win?",
        f"Should the team accept or mitigate the principal capability gap: {_first_sentences(rating.get('per_field_reasoning', {}).get('capability_fit'), 1, 220)}",
    ]
    if context["submission_deadline"] != MISSING:
        asks.append(f"If GO, should capture activity start now against the {context['submission_deadline']} deadline?")
    asks.append("Should this proceed solo, with a delivery partner, or be parked pending stronger evidence?")
    return asks


def _executive_summary(context):
    ai_summary = context["executive_summary_ai"]
    if ai_summary.get("opening") and ai_summary.get("scope_summary") and ai_summary.get("executive_view"):
        return ai_summary["opening"], ai_summary["scope_summary"], ai_summary["executive_view"]
    verb = "is seeking market input on" if context["uk_stage"] == "UK2" else "is procuring"
    scope = _scope_points(context)
    first = f"{context['buyer']} {verb} {context['title']}."
    second = "Published scope: " + " ".join(scope[:3])
    reasoning = context["ai_read"].get("per_field_reasoning", {})
    third = (
        f"Executive view: {context['ai_read']['overall']} (provisional). Capability fit is "
        f"{context['ai_read']['capability_fit']} and right to win is {context['ai_read']['right_to_win']}. "
        f"{_first_sentences(reasoning.get('overall'), 2, 420)}"
    )
    return first, second, third


def _save(document, output_dir, filename):
    output = Path(output_dir) / filename
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    return str(output)


def _set_footer(document, context):
    for section in document.sections:
        paragraph = section.footer.paragraphs[0]
        _set_paragraph(
            paragraph,
            f"Smarter Bids. Real Results. | © 2026 Bid Savvy Solutions Ltd | "
            f"{context['notice_reference']} | {context['generated_at'][:19].replace('T', ' ')}",
        )


_SECTION_D_VERBATIM = (
    "Listed below are only genuine blockers, meaning things that would physically prevent or "
    "materially complicate a bid: an unresolved type of work, an absent product, a named "
    "framework, or a hard deadline. Per Victoria's direction of 11 August 2026, Trifork's UK "
    "track record, entity size, buyer relationships and clearance position are not treated as "
    "blockers or risks. Establishing UK delivery references is the purpose of the Bid Savvy "
    "engagement, and those points are handled in capture strategy, not flagged as reasons for "
    "caution."
)


def build_internal_addendum_docx(conn, notice_id, output_dir):
    context = build_context(conn, notice_id)
    document = Document(TEMPLATE_DIR / "internal_addendum_template.docx")
    tables = document.tables

    # Section C's heading must name the actual rating; it was previously a
    # static "HIGH" regardless of the notice's real capability_fit.
    _set_paragraph(document.paragraphs[6], f"C. WHY THIS IS A {context['ai_read']['capability_fit']} FIT")

    _set_cell(
        tables[0].cell(0, 0),
        f"INTERNAL ADDENDUM | NOT FOR CLIENT DISTRIBUTION\n{context['title']}\n"
        f"{context['buyer']} | Reference: {context['notice_reference']} | {_urgency_text(context)}\n"
        f"Prepared by {context['owner_name']}, Bid Savvy Solutions Ltd | {context['generated_at'][:10]}",
    )
    _set_cell(
        tables[1].cell(0, 0),
        "INTERNAL USE ONLY — NOT FOR CLIENT DISTRIBUTION. This document is for Victoria Milan's "
        "GO / NO-GO / Park decision and must not be shared externally.",
    )
    _shade_cell(tables[1].cell(0, 0), **_CALLOUT_RED)
    gates = context["gate_outcomes"][:5]
    for index in range(1, 6):
        gate = gates[index - 1] if index <= len(gates) else None
        _set_cell(tables[2].cell(index, 0), f"Gate {index}: {gate['gate_name']}" if gate else f"Gate {index}")
        _set_cell(tables[2].cell(index, 1), f"{gate['result']}. {gate['reason']}" if gate else MISSING)

    metadata = (
        ("Spotted by", f"{context['owner_name']}, Bid Savvy Solutions Ltd" if context["owner_name"] != MISSING else MISSING),
        ("Date spotted", context["date_spotted"]),
        ("Pipeline reference", _pipeline_reference(context)),
        ("Phase 2 status", f"{context['ai_read']['overall']} — PROVISIONAL, FOR VALIDATION"),
        ("Client brief status", _client_brief_status_text(context)),
        ("Tracker status", _tracker_status_text(context)),
        ("Notice reference", _notice_reference_with_source(context)),
    )
    for row_index, (field, value) in enumerate(metadata, start=1):
        _set_cell(tables[3].cell(row_index, 0), field)
        if field == "Notice reference" and context["notice_url"] != MISSING:
            _set_hyperlink(tables[3].cell(row_index, 1), value, context["notice_url"])
        else:
            _set_cell(tables[3].cell(row_index, 1), value)

    _insert_med_dual_reading(document, tables[4]._tbl, context)
    mapping_rows = _capability_mapping_rows(context)
    _set_cell(tables[4].cell(0, 0), "Buyer problem")
    _set_cell(tables[4].cell(0, 1), "Trifork capability mapping")
    for row_index in range(1, min(len(tables[4].rows), len(mapping_rows) + 1)):
        problem, capability = mapping_rows[row_index - 1]
        _set_cell(tables[4].cell(row_index, 0), problem)
        _set_cell(tables[4].cell(row_index, 1), capability)
    _trim_table_rows(tables[4], max(1, len(mapping_rows)))

    _insert_paragraph_before(document, tables[5]._tbl, _SECTION_D_VERBATIM)
    risks = context["blockers_risks"]
    if risks:
        for row_index in range(1, min(len(tables[5].rows), len(risks) + 1)):
            blocker, assessment = _split_blocker(risks[row_index - 1])
            _set_cell(tables[5].cell(row_index, 0), blocker)
            _set_cell(tables[5].cell(row_index, 1), assessment)
    else:
        # House style (2026-08-19): an empty blockers list is a valid,
        # genuine result, and must read as one clear line, not a padded or
        # blank-dashed table. Removing the table element entirely would
        # shift every later table's position once the file is reopened
        # (breaks the fixed-index reads elsewhere), so this keeps one row
        # carrying the exact required line instead.
        _set_cell(tables[5].cell(1, 0), "No genuine blockers identified.")
        _set_cell(tables[5].cell(1, 1), "")
    _trim_table_rows(tables[5], max(1, len(risks)))

    asks = context["direct_asks"] or _victoria_asks(context)
    for row_index in range(1, min(len(tables[6].rows), len(asks) + 1)):
        ask_text, why_text = _split_ask(asks[row_index - 1])
        _set_cell(tables[6].cell(row_index, 0), ask_text)
        _set_cell(tables[6].cell(row_index, 1), why_text)
    _trim_table_rows(tables[6], max(1, len(asks)))

    _set_cell(tables[7].cell(0, 0), _final_decision_text(context))
    _shade_cell(tables[7].cell(0, 0), **_CALLOUT_BLUE)
    _set_paragraph(
        document.paragraphs[-1],
        f"Prepared by {context['owner_name']}, Bid Savvy Solutions Ltd | Internal use only | "
        f"Not for client distribution | {context['generated_at'][:10]}",
    )
    _set_footer(document, context)
    safe_ref = context["notice_reference"].replace("/", "-")
    return _save(document, output_dir, f"{safe_ref}_internal_addendum.docx")


def build_capture_brief_docx(conn, notice_id, output_dir):
    context = build_context(conn, notice_id)
    document = Document(TEMPLATE_DIR / "capture_brief_template.docx")
    tables = document.tables
    paragraphs = document.paragraphs

    _set_cell(
        tables[0].cell(0, 0),
        f"CAPTURE BRIEF\n{context['title']}\n{context['buyer']} | Reference: "
        f"{context['notice_reference']} | {_urgency_text(context)}\nPrepared for Victoria Milan | "
        f"{context['generated_at'][:10]}",
    )
    executive_view = context["executive_summary_ai"].get("executive_view") or _first_sentences(
        context["ai_read"].get("per_field_reasoning", {}).get("capability_fit"), 2, 420
    )
    _set_cell(
        tables[1].cell(0, 0),
        f"Why this matters: {executive_view} "
        f"Decision point: {_recommended_action(context)} PROVISIONAL, FOR VALIDATION.",
    )
    _shade_cell(tables[1].cell(0, 0), **_CALLOUT_BLUE)
    _set_cell(tables[2].cell(0, 0), f"Submission deadline: {context['submission_deadline']} | {_urgency_text(context)}")
    _shade_cell(tables[2].cell(0, 0), **_CALLOUT_RED)

    key_terms = _key_terms_rows(context) or [("None", "The notice is in plain English; no jargon requires explanation.")]
    for row_index in range(1, min(len(tables[3].rows), len(key_terms) + 1)):
        term, meaning = key_terms[row_index - 1]
        _set_cell(tables[3].cell(row_index, 0), term)
        _set_cell(tables[3].cell(row_index, 1), meaning)
    _trim_table_rows(tables[3], len(key_terms))

    info = (
        ("Contracting authority", context["buyer"]),
        ("Reference", context["notice_reference"]),
        ("Notice type", _notice_type_label(context)),
        ("Route to market", _route_to_market_text(context)),
        ("Framework status", context["framework_status"]),
        ("Estimated contract value", _format_value(context)),
        ("Main CPV codes", ", ".join(context["cpv_codes"]) or MISSING),
        ("Contact", context["buyer_contact"]),
        ("Notice link", "Open published notice"),
        ("Owner", context["owner_name"]),
    )
    _ensure_table_rows(tables[4], len(info) + 1)
    for row_index, (field, value) in enumerate(info, start=1):
        _set_cell(tables[4].cell(row_index, 0), field)
        if field == "Notice link" and context["notice_url"] != MISSING:
            _set_hyperlink(tables[4].cell(row_index, 1), value, context["notice_url"])
        else:
            _set_cell(tables[4].cell(row_index, 1), value)

    milestones = _timetable_rows(context)
    for row_index in range(1, min(len(tables[5].rows), len(milestones) + 1)):
        field, value = milestones[row_index - 1]
        _set_cell(tables[5].cell(row_index, 0), field)
        _set_cell(tables[5].cell(row_index, 1), value)
    _trim_table_rows(tables[5], max(1, len(milestones)))

    for column, rating in enumerate((
        context["ai_read"]["capability_fit"],
        context["ai_read"]["competitor_position"],
        context["ai_read"]["right_to_win"],
    )):
        cell = tables[6].cell(1, column)
        _set_cell(cell, rating)
        _color_rating_run(cell.paragraphs[0], rating)
        cell.paragraphs[0].alignment = 1  # WD_ALIGN_PARAGRAPH.CENTER

    decision_rows = _decision_framework_rows(context)
    for row_index in range(1, min(len(tables[7].rows), len(decision_rows) + 1)):
        question, implication = decision_rows[row_index - 1]
        _set_cell(tables[7].cell(row_index, 0), question)
        _set_cell(tables[7].cell(row_index, 1), implication)
    _trim_table_rows(tables[7], max(1, len(decision_rows)))

    action_rows = _immediate_action_rows(context)
    for row_index in range(1, min(len(tables[8].rows), len(action_rows) + 1)):
        action, owner = action_rows[row_index - 1]
        _set_cell(tables[8].cell(row_index, 0), action)
        _set_cell(tables[8].cell(row_index, 1), owner)
    _trim_table_rows(tables[8], max(1, len(action_rows)))

    summary = (
        ("Opportunity", context["title"]), ("Buyer", context["buyer"]),
        ("Reference", context["notice_reference"]),
        ("Estimated value", _format_value(context)),
        ("Route to market", _route_to_market_text(context)),
        ("Capability fit", context["ai_read"]["capability_fit"]),
        ("Competitive risk", context["ai_read"]["competitor_position"]),
        ("Right to win", context["ai_read"]["right_to_win"]),
        ("Recommended action", _recommended_action(context)),
        ("Next deadline", context["submission_deadline"]),
        ("Sources", _sources_text(context)),
    )
    for row_index, (field, value) in enumerate(summary, start=1):
        _set_cell(tables[9].cell(row_index, 0), field)
        _set_cell(tables[9].cell(row_index, 1), value)

    executive = _executive_summary(context)
    requirement_areas = _requirement_areas(context)
    replacements = {
        2: executive[0],
        3: executive[1],
        4: executive[2] + " PROVISIONAL, FOR VALIDATION.",
        14: "The published notice describes the following scope:",
        23: "What the buyer is seeking from this engagement",
        24: _what_buyer_seeking(context),
        25: "Engagement model",
        26: _engagement_model_text(context),
        30: f"Capability fit: {context['ai_read']['capability_fit']}",
        31: context["ai_read"].get("per_field_reasoning", {}).get("capability_fit", MISSING),
        36: f"Competitive risk: {context['ai_read']['competitor_position']}",
        37: context["ai_read"].get("per_field_reasoning", {}).get("competitor_position", MISSING),
        38: f"Right to win: {context['ai_read']['right_to_win']}",
        39: context["ai_read"].get("per_field_reasoning", {}).get("right_to_win", MISSING),
        46: _solo_or_partner_text(context)[0],
        47: _solo_or_partner_text(context)[1],
        51: f"Prepared by {context['owner_name']}, Bid Savvy Solutions Ltd",
        52: "Confidential | Not for circulation beyond Trifork UK",
    }
    for index, value in replacements.items():
        if index < len(paragraphs):
            _set_paragraph(paragraphs[index], value)
    # paragraphs[15] is the "Documented requirement areas" sub-heading itself
    # and must be left untouched; only 16-22 (7 slots) are the bullet items.
    scope_paragraphs = paragraphs[16:23]
    for index, paragraph in enumerate(scope_paragraphs):
        if index < len(requirement_areas):
            _set_paragraph(paragraph, requirement_areas[index])
        else:
            _remove_paragraph(paragraph)
    # paragraphs[32:36] are the sample's dedicated evidence slots -- one named
    # Trifork case study per paragraph, backing up the capability_fit claim
    # in paragraph 31 with something concrete Victoria can check. Previously
    # these were unconditionally deleted instead of populated, which is why
    # Section 6 read as one generic sentence instead of the sample's cited
    # case-study list.
    # Paragraph 31 already carries the single capability_fit reasoning
    # sentence, so the dual-reading insert here skips re-adding it and only
    # contributes the build/product signal lists plus the honest-position
    # line when the rating is genuinely MED.
    _insert_med_dual_reading(document, paragraphs[32]._element, context, include_reasoning=False)

    case_study_paragraphs = paragraphs[32:36]
    case_studies = [row.get("capability_mapping", MISSING) for row in context["capability_mapping"]][:4]
    for index, paragraph in enumerate(case_study_paragraphs):
        if index < len(case_studies):
            _set_paragraph(paragraph, case_studies[index])
        else:
            _remove_paragraph(paragraph)

    _set_footer(document, context)

    safe_ref = context["notice_reference"].replace("/", "-")
    return _save(document, output_dir, f"{safe_ref}_capture_brief.docx")
